"""
Generate Trading_Agent_B_Design.docx from Strategy B DESIGN.md content.
Run: python3 generate_design_b.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERSION = "v1.2"
DATE    = "2026-05-22"

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
ORANGE = RGBColor(0xF4, 0x7B, 0x20)
SLATE  = RGBColor(0x44, 0x55, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF2, 0xF5, 0xF9)


def _shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def _header_row(table, *labels):
    row = table.rows[0]
    for i, label in enumerate(labels):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = label
        _shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9)


def _data_row(table, row_idx, *values, shade=False):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = str(val)
        if shade:
            _shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = SLATE


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else SLATE
    return p


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color or SLATE
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = SLATE


def title_block(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trading Agent B")
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("System Design Document")
    run2.font.size  = Pt(14)
    run2.font.color.rgb = ORANGE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{VERSION}  ·  {DATE}  ·  Alpaca Paper Trading  ·  Blue Chip Pool Strategy")
    run3.font.size   = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = SLATE

    doc.add_paragraph()


def section_overview(doc):
    add_heading(doc, "1. What It Is")
    add_body(doc,
        "An autonomous day-trading system that operates on a curated blue chip universe using a "
        "3-tier behavioral pool system. Rather than scanning 600+ tickers broadly, Strategy B "
        "maintains a focused shortlist of stocks that have demonstrated consistent behavioral "
        "patterns — predictable VWAP respect, reliable ATR ranges, and repeatable momentum setups.")
    add_body(doc, "Daily objective: $300–$700 realized P&L via high-conviction position management on proven performers.", bold=True)
    add_body(doc,
        "Core differentiator: Every stock Claude sees has been pre-qualified through behavioral scoring. "
        "Claude receives rolling_score, above_vwap, rs_vs_sector, and atr_ratio context per stock, "
        "enabling more precise confidence assignment than broad scanning allows.")
    doc.add_paragraph()


def section_architecture(doc):
    add_heading(doc, "2. Architecture")

    add_code(doc,
        "cron-job.org  →  GitHub Actions (trading-agent-b)\n"
        "                         │\n"
        "                         ▼\n"
        "              orchestrator.py\n"
        "                    │\n"
        "   ┌────────────────┼────────────────────────┐\n"
        "   ▼                ▼                        ▼\n"
        "premarket()    intraday()                 eod()\n"
        "pool_filter    position mgmt             close_all\n"
        "scanner        momentum scan             pool_scorer\n"
        "strategy       strategy+risk             daily_perf\n"
        "risk+guards    guards+execute\n"
        "execute\n"
        "   │\n"
        "   ▼\n"
        "Supabase (b_ tables)  +  Alpaca Paper (strategy=b tag)")

    add_heading(doc, "Stack", level=2)
    items = [
        "Python 3.11 · Claude claude-opus-4-7 (Anthropic) — higher quality for blue chip behavioral context",
        "Alpaca Markets API — paper trading, bracket orders, live quotes; all orders tagged strategy=b",
        "Supabase (PostgreSQL) — b_ prefixed tables in same project as Strategy A",
        "Streamlit Cloud — Strategy B dashboard tab + combined A vs B view",
        "GitHub Actions — separate repo from Strategy A; independent deployment and failure modes",
        "cron-job.org — external trigger; intraday every 30 min (vs 15 min for Strategy A)",
    ]
    for item in items:
        add_bullet(doc, item)
    doc.add_paragraph()


def section_pool_system(doc):
    add_heading(doc, "3. Pool System")
    add_body(doc,
        "The 3-tier pool is Strategy B's core differentiator. Every stock must earn its place "
        "through behavioral scoring before Claude ever sees it.")
    doc.add_paragraph()

    add_heading(doc, "3.1  Pool 1 — Full Blue Chip Universe (~150 stocks)", level=2)
    add_body(doc,
        "The complete eligible universe: large-cap, liquid, institutionally-owned stocks. "
        "Market cap > $10B, avg volume > 2M shares/day. Seeded from POOL_2_SEED config. "
        "Stocks demoted from Pool 2 return here.")
    doc.add_paragraph()

    add_heading(doc, "3.2  Pool 2 — Behavioral Shortlist (~25–50 stocks)", level=2)
    add_body(doc,
        "Stocks with consistent, repeatable behavioral patterns for this strategy over rolling "
        "7-day windows. This is the active working set — Pool Filter selects Pool 3 exclusively "
        "from Pool 2.")
    for item in [
        "Consistent VWAP respect (price bounces off VWAP with low slippage)",
        "ATR moves that align with the +2% target without excessive overshoot",
        "Volume patterns that confirm momentum (surge at breakout)",
        "Sector relative strength that supports directional bias",
    ]:
        add_bullet(doc, item)
    add_body(doc,
        "Promotion/demotion: Pool Scorer runs EOD. Stocks with rolling_score >= 6 are promoted to "
        "Pool 2; stocks with rolling_score < 3 are demoted to Pool 1.")
    doc.add_paragraph()

    add_heading(doc, "3.3  Pool 3 — Daily Elite Picks (8–10 stocks)", level=2)
    add_body(doc,
        "Selected fresh each morning by Pool Filter from Pool 2. These are the only stocks Claude "
        "can trade on any given day. Pool 3 is ephemeral — re-derived each morning from Pool 2 "
        "using 5 real-time signals:")

    table = doc.add_table(rows=6, cols=2)
    _style_table(table)
    _set_col_widths(table, [5, 11.5])
    _header_row(table, "Signal", "What It Measures")
    signals = [
        ("vwap_position",       "Is price above VWAP at market open? High weight — primary filter."),
        ("orb_breakout",        "Did price break above Opening Range (9:30–9:45 AM high)?"),
        ("volume_ratio",        "Current volume vs 20-day average. Surge = bullish confirmation."),
        ("sector_rs",           "Sector relative strength vs SPY for the day."),
        ("volume_acceleration", "Rate of volume build in first 30 minutes."),
    ]
    for i, (sig, desc) in enumerate(signals):
        _data_row(table, i + 1, sig, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
    doc.add_paragraph()


def section_agent_inventory(doc):
    add_heading(doc, "4. Agent Inventory")

    agents = [
        ("Pool Manager",               "core/pool_manager.py",
         "POOL_2_SEED config, b_pools",
         "Updated b_pools",
         "Maintains Pool 1/2/3 membership. Seeds from config if empty. Called at startup."),
        ("Pool Filter",                "scanner/pool_filter.py",
         "Pool 2 stocks, real-time data",
         "Pool 3 candidates (8-10) with scores",
         "Each morning: fetches live data for all Pool 2 stocks, scores on 5 signals, returns top picks."),
        ("Scanner",                    "scanner/scanner.py",
         "Pool 3 candidates",
         "Behavioral scored candidates",
         "VWAP respect, ATR alignment, volume patterns, RSI, MACD, SMA. Returns behavioral_score."),
        ("News Intel",                 "agents/news_intel.py",
         "Scored candidates",
         "Filtered candidates + news context",
         "Earnings blackout filter. Removes earnings-day tickers. Adds news sentiment."),
        ("Market Context",             "agents/market_context.py",
         "VIX, Fear & Greed, futures",
         "market_context dict, quiet_day flag",
         "VIX/F&G/futures gate. Hard skip if futures < -1.5%. Sets max_positions."),
        ("Strategy Agent",             "agents/strategy.py",
         "Enriched candidates, per-stock behavioral data",
         "Trade plan (entry/target/stop/confidence/reasoning)",
         "Claude claude-opus-4-7. Reads pool, rolling_score, above_vwap, rs_vs_sector, atr_ratio, signal_type. INTRADAY_MOMENTUM bypasses 1pm restriction. Prompt cached."),
        ("Risk Agent",                 "agents/risk.py",
         "Trade plan",
         "Validated/rejected trades",
         "R:R >= 2.0, position size $2500-$3500, stop <= 0.67%, max loss check."),
        ("Sector Guard",               "agents/sector_guard.py",
         "Validated trades, open positions",
         "Filtered trades",
         "Sector concentration cap. SECTOR_MAP + yfinance fallback."),
        ("Guardrails",                 "agents/guardrails.py",
         "Trade list, daily P&L state",
         "Final approved trades",
         "Duplicates, price sanity (>5% from market = reject), daily loss limit, capital cap."),
        ("Pool Scorer",                "agents/pool_scorer.py",
         "Today's Pool 3 trades",
         "Updated b_stock_scores, updated b_pools",
         "EOD: win/loss, P&L, slippage, setup alignment. 7-day rolling score. Promotes (>=6) or demotes (<3)."),
        ("Intraday Momentum Scanner",  "scanner/intraday_momentum.py",
         "Pool 3 stocks, SPY snapshot, live prices",
         "Momentum candidates (pool=2, signal_type=INTRADAY_MOMENTUM)",
         "SPY gate (>=+0.5%). Scans Pool 3 for stocks up >=0.5% above VWAP. Max 6 runs/day, 90 min interval."),
        ("Alpaca Broker",              "agents/alpaca_broker.py",
         "Validated trades",
         "Bracket order confirmations, live data",
         "Alpaca API wrapper. Bracket orders tagged strategy=b. get_orders, snapshots, live prices."),
    ]

    table = doc.add_table(rows=len(agents) + 1, cols=5)
    _style_table(table)
    _set_col_widths(table, [3.5, 4.0, 3.5, 3.5, 5.5])
    _header_row(table, "Agent", "File", "Inputs", "Outputs", "Responsibility")
    for i, (agent, file_, inputs, outputs, resp) in enumerate(agents):
        _data_row(table, i + 1, agent, file_, inputs, outputs, resp, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()


def section_pipeline(doc):
    add_heading(doc, "5. Daily Pipeline")

    # Premarket
    add_heading(doc, "5.1  Premarket — 10:00 AM ET", level=2)
    add_body(doc,
        "Runs once before significant intraday volume develops. Pool Filter derives Pool 3 first, "
        "then the full pipeline runs on those 8–10 stocks only.")
    doc.add_paragraph()

    table = doc.add_table(rows=11, cols=3)
    _style_table(table)
    _set_col_widths(table, [3, 4, 9.5])
    _header_row(table, "Step", "Agent", "What Happens")
    steps = [
        ("0. Pool Setup",        "Pool Manager",   "Verifies pool membership. Seeds from config if b_pools is empty."),
        ("1. Pool Filter",       "Pool Filter",    "Scores all Pool 2 stocks on 5 real-time signals. Returns top 8–10 as Pool 3."),
        ("2. Behavioral Scan",   "Scanner",        "Scores Pool 3: VWAP respect, ATR alignment, volume patterns, RSI, MACD, SMA20/50."),
        ("3. News Filter",       "News Intel",     "Removes earnings-day tickers. Adds news sentiment context."),
        ("4. Market Context",    "Market Context", "VIX, Fear & Greed, futures. Sets max_positions. Hard skip if futures < -1.5%."),
        ("5. Strategy (Claude)", "Strategy Agent", "claude-opus-4-7 selects trades with full behavioral context per stock."),
        ("6. Risk Validation",   "Risk Agent",     "R:R >= 2.0, size bounds, stop validation."),
        ("7. Sector Guard",      "Sector Guard",   "Caps sector concentration."),
        ("8. Guardrails",        "Guardrails",     "Duplicates, price sanity, daily loss limit."),
        ("9. Execute",           "Alpaca Broker",  "Bracket orders tagged strategy=b. Records to b_positions."),
    ]
    for i, (step, agent, desc) in enumerate(steps):
        _data_row(table, i + 1, step, agent, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()

    # Intraday
    add_heading(doc, "5.2  Intraday — Every 30 min, 10:00 AM–3:45 PM ET", level=2)
    add_body(doc, "Every cycle (position management):")
    for item in [
        "Reconcile: detect bracket exits (stop/target). Record exit price and P&L.",
        "Refresh: sync current price and unrealized P&L.",
        "Lock-in: Tier 1 ($500 realized) — tighten trail. Tier 2 ($700 total) — close all.",
    ]:
        add_bullet(doc, item)

    add_body(doc, "Momentum scan (conditional, max 6/day, min 90 min apart):")
    for item in [
        "Guards: check max runs, min interval, open slots, loss limit — skip if any fail.",
        "SPY gate: SPY must be up >= +0.5% at scan time.",
        "Scan Pool 3 for stocks up >= +0.5% above VWAP.",
        "Candidates tagged pool=2, signal_type=INTRADAY_MOMENTUM.",
        "Run through Strategy → Risk → Guardrails → Execute. Target capped at +1%.",
        "INTRADAY_MOMENTUM signal bypasses 1pm pool restriction.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # EOD
    add_heading(doc, "5.3  EOD — Post 4:00 PM ET", level=2)
    table = doc.add_table(rows=5, cols=2)
    _style_table(table)
    _set_col_widths(table, [4.5, 12])
    _header_row(table, "Step", "What Happens")
    eod_steps = [
        ("Close Positions",   "Market-sell all open positions. Cancel pending bracket legs."),
        ("Pool Scoring",      "Score each Pool 3 stock: win/loss, P&L, slippage vs ATR, setup alignment. "
                              "Write to b_stock_scores. Compute 7-day rolling_score. Promote (>=6) or demote (<3) between pools."),
        ("Daily Performance", "Write P&L summary to b_daily_performance: realized, win rate, count, best/worst trade."),
        ("Daily Summary",     "Generate narrative summary for review."),
    ]
    for i, (step, desc) in enumerate(eod_steps):
        _data_row(table, i + 1, step, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()


def section_trading_logic(doc):
    add_heading(doc, "6. Trading Logic")

    # Position sizing
    add_heading(doc, "6.1  Position Sizing", level=2)
    table = doc.add_table(rows=4, cols=3)
    _style_table(table)
    _set_col_widths(table, [3, 2.5, 11])
    _header_row(table, "Confidence", "Size", "Trigger Criteria")
    rows = [
        ("HIGH",   "$3,500", "rolling_score >= 7  AND  above_vwap = True  AND  rs_vs_sector >= 1.5"),
        ("MEDIUM", "$3,000", "rolling_score 4-6  OR  (rolling_score 3-4  AND  above_vwap  AND  rs_vs_sector >= 1.2)"),
        ("LOW",    "$2,500", "rolling_score 3-4 with weaker VWAP/RS signals"),
    ]
    for i, (conf, size, trigger) in enumerate(rows):
        _data_row(table, i + 1, conf, size, trigger, shade=(i % 2 == 1))
    doc.add_paragraph()

    # Trade formulas
    add_heading(doc, "6.2  Trade Formulas", level=2)
    add_body(doc, "All values set by fixed formula — Claude cannot deviate:")
    add_code(doc,
        "entry_price    = Alpaca ask price (live) or scanner close\n"
        "target_price   = round(entry × 1.0200, 2)   # +2.0% full target (premarket)\n"
        "partial_target = round(entry × 1.0100, 2)   # +1.0% partial exit (Leg A)\n"
        "stop_loss      = round(entry × 0.9933, 2)   # -0.67% stop\n"
        "shares         = int(position_size / entry)\n"
        "\n"
        "intraday target = round(entry × 1.0100, 2)  # +1% only — shorter window\n"
        "\n"
        "Reward:Risk (premarket) = 2.00% / 0.67% = 2.99 ≈ 3:1\n"
        "R:R floor               = 2.0 (both premarket and intraday)"
    )
    doc.add_paragraph()

    # Partial profit design
    add_heading(doc, "6.3  Partial Profit Design", level=2)
    add_body(doc, "Each premarket trade opens as two independent bracket orders:")
    add_code(doc,
        "Leg A  →  shares // 2    ·  target = entry × 1.01  (+1%)\n"
        "Leg B  →  shares - A     ·  target = entry × 1.02  (+2%)\n"
        "Both   →  stop_loss = entry × 0.9933"
    )
    doc.add_paragraph()

    # Time-of-day rules
    add_heading(doc, "6.4  Time-of-Day Rules", level=2)
    table = doc.add_table(rows=5, cols=3)
    _style_table(table)
    _set_col_widths(table, [4, 5.5, 7])
    _header_row(table, "Time", "Rule", "Rationale")
    tod_rows = [
        ("10:00–12:59 PM", "New entries from Pool 2/3 stocks allowed", "Full day ahead"),
        ("1:00–3:45 PM",   "Pool 3 only, no new Pool 2 entries",       "Less time to recover from bad entry"),
        ("3:45 PM+",       "No new entries, manage existing only",     "Too close to close"),
        ("INTRADAY_MOMENTUM (any hour)",
         "Exempt from 1pm pool restriction",
         "Confirmed movers — SPY gate + VWAP confirmation already fired"),
    ]
    for i, (time_, rule, reason) in enumerate(tod_rows):
        _data_row(table, i + 1, time_, rule, reason, shade=(i % 2 == 1))
    doc.add_paragraph()

    # Trailing stop
    add_heading(doc, "6.5  Trailing Stop", level=2)
    add_code(doc,
        "effective_stop = max(stop_loss, high_watermark × (1 - 1.0%))\n"
        "\n"
        "After Tier 1 lock-in ($500 realized):\n"
        "effective_stop = max(stop_loss, high_watermark × (1 - 0.5%))"
    )
    doc.add_paragraph()


def section_risk(doc):
    add_heading(doc, "7. Risk Controls")
    add_body(doc, "Five independent layers applied in sequence — any one can block a trade:")
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=3)
    _style_table(table)
    _set_col_widths(table, [3.5, 4, 9])
    _header_row(table, "Layer", "Agent", "What It Blocks")
    guards = [
        ("Market Gate",         "market_context.py", "Trading on crash days (futures < -1.5%), extreme volatility"),
        ("News Filter",         "news_intel.py",     "Earnings-day tickers, negative catalyst stocks"),
        ("Risk Agent",          "risk.py",           "R:R below 2.0 floor, position size out of bounds, stop too wide"),
        ("Sector Guard",        "sector_guard.py",   "Sector concentration breaches"),
        ("Guardrails",          "guardrails.py",     "Duplicates, price sanity >5% from market, daily loss limit -$250"),
    ]
    for i, (layer, file_, desc) in enumerate(guards):
        _data_row(table, i + 1, layer, file_, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()

    add_heading(doc, "7.1  Intraday Momentum Guards", level=2)
    table = doc.add_table(rows=6, cols=3)
    _style_table(table)
    _set_col_widths(table, [4, 3.5, 9])
    _header_row(table, "Guard", "Threshold", "Reason")
    guards2 = [
        ("Max daily runs",    "6",        "Prevents overtrading on volatile days"),
        ("Min interval",      "90 min",   "Ensures genuine new signal, not noise"),
        ("SPY gate",          ">=+0.5%",  "Confirms broad market support for momentum"),
        ("Open slots",        "< max_pos","Must have capacity before scanning"),
        ("Loss limit",        "> -$250",  "Don't compound losses with momentum trades"),
    ]
    for i, (guard, thresh, reason) in enumerate(guards2):
        _data_row(table, i + 1, guard, thresh, reason, shade=(i % 2 == 1))
    doc.add_paragraph()

    add_heading(doc, "7.2  Pool Behavioral Gate", level=2)
    add_body(doc,
        "The Pool Filter itself is a risk layer. Only stocks with proven behavioral track records "
        "enter Pool 3. A stock that consistently stops out, shows wide slippage, or fails to hit "
        "targets gets demoted via the Pool Scorer — pre-screening away setups that are technically "
        "valid but behaviorally unreliable before they ever reach Claude.")
    doc.add_paragraph()


def section_config(doc):
    add_heading(doc, "8. Key Configuration")

    params = [
        ("TOTAL_CAPITAL",               "$50,000",  "Simulated account size"),
        ("TARGET_PCT",                  "2.0%",     "Full profit target (premarket)"),
        ("INTRADAY_TARGET_PCT",         "1.0%",     "Intraday entry profit target"),
        ("PARTIAL_PROFIT_PCT",          "1.0%",     "Partial exit (Leg A)"),
        ("MAX_LOSS_PER_TRADE",          "0.67%",    "Stop loss depth"),
        ("MIN_REWARD_RISK",             "2.0",      "R:R floor (both premarket and intraday)"),
        ("TRAIL_PCT",                   "1.0%",     "Trailing stop from high watermark"),
        ("LOCK_IN_TRAIL_PCT",           "0.5%",     "Tighter trail after Tier 1"),
        ("DAILY_LOCK_IN_TARGET",        "$500",     "Tier 1: let winners ride"),
        ("DAILY_BONUS_TARGET",          "$700",     "Tier 2: close everything"),
        ("DAILY_LOSS_LIMIT",            "-$250",    "Stop trading (0.5% of capital)"),
        ("MAX_POSITIONS",               "10",       "Max concurrent positions"),
        ("MAX_INTRADAY_RUNS",           "6",        "Max momentum scan runs per day"),
        ("MIN_INTRADAY_INTERVAL_MIN",   "90",       "Min minutes between momentum scans"),
        ("SPY_MOMENTUM_GATE",           "+0.5%",    "SPY min gain to allow momentum scan"),
        ("POOL_PROMOTE_THRESHOLD",      "6.0",      "7-day rolling score to promote to Pool 2"),
        ("POOL_DEMOTE_THRESHOLD",       "3.0",      "7-day rolling score to demote from Pool 2"),
        ("HIGH_CONFIDENCE_SIZE",        "$3,500",   "HIGH confidence position size"),
        ("MEDIUM_CONFIDENCE_SIZE",      "$3,000",   "MEDIUM confidence position size"),
        ("LOW_CONFIDENCE_SIZE",         "$2,500",   "LOW confidence position size"),
    ]

    table = doc.add_table(rows=len(params) + 1, cols=3)
    _style_table(table)
    _set_col_widths(table, [6, 2.5, 8])
    _header_row(table, "Parameter", "Value", "Purpose")
    for i, (param, val, desc) in enumerate(params):
        _data_row(table, i + 1, param, val, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
    doc.add_paragraph()


def section_data_model(doc):
    add_heading(doc, "9. Data Model — Supabase Tables")
    add_body(doc, "All tables use b_ prefix to share the same Supabase project as Strategy A.")
    doc.add_paragraph()

    tables_info = [
        ("b_pools",
         [("id", "uuid", "primary key"),
          ("ticker", "text", "stock symbol"),
          ("pool", "int", "1, 2, or 3"),
          ("added_date", "date", "when stock entered this pool"),
          ("rolling_score", "numeric", "current 7-day weighted average"),
          ("last_traded", "date", "last date this stock was traded"),
          ("notes", "text", "manual annotation")]),
        ("b_stock_scores",
         [("id", "uuid", "primary key"),
          ("date", "date", "scoring date"),
          ("ticker", "text", "stock symbol"),
          ("pool", "int", "pool at time of scoring"),
          ("traded", "boolean", "was it actually traded today?"),
          ("win", "boolean", "null if not traded"),
          ("pnl", "numeric", "null if not traded"),
          ("slippage_bps", "numeric", "realized slippage in basis points"),
          ("setup_score", "numeric", "0-10, scanner signal alignment"),
          ("daily_score", "numeric", "weighted composite for this day"),
          ("rolling_7d", "numeric", "7-day weighted average")]),
        ("b_positions",
         [("id", "uuid", "primary key"),
          ("date", "date", "trading date"),
          ("run_id", "uuid", "FK to b_daily_runs"),
          ("ticker", "text", "stock symbol"),
          ("status", "text", "OPEN / CLOSED / UNFILLED / STOP / TARGET"),
          ("pool", "int", "pool when entered"),
          ("signal_type", "text", "PREMARKET / INTRADAY_MOMENTUM"),
          ("entry_price", "numeric", "filled entry price"),
          ("exit_price", "numeric", "filled exit price"),
          ("realized_pnl", "numeric", "closed P&L in $"),
          ("close_reason", "text", "TARGET / STOP / TRAIL / LOCK_IN / EOD"),
          ("alpaca_order_id", "text", "Alpaca order reference")]),
        ("b_daily_runs",
         [("id", "uuid", "primary key"),
          ("date", "date", "trading date"),
          ("run_number", "int", "0=premarket, 1-6=intraday momentum"),
          ("run_type", "text", "PREMARKET / INTRADAY_MOMENTUM"),
          ("started_at", "timestamptz", "run start time"),
          ("pool3_tickers", "text[]", "Pool 3 stocks for this run"),
          ("trades_placed", "int", "number of orders placed"),
          ("skipped_reason", "text", "null if run completed normally")]),
        ("b_daily_performance",
         [("id", "uuid", "primary key"),
          ("date", "date", "trading date"),
          ("realized_pnl", "numeric", "total realized P&L"),
          ("win_rate", "numeric", "fraction of winning trades"),
          ("positions_count", "int", "number of trades taken"),
          ("momentum_pnl", "numeric", "P&L from INTRADAY_MOMENTUM entries only"),
          ("pool2_pnl", "numeric", "P&L from Pool 2 stocks"),
          ("pool3_pnl", "numeric", "P&L from Pool 3 stocks")]),
    ]

    for tname, cols in tables_info:
        add_heading(doc, tname, level=2)
        table = doc.add_table(rows=len(cols) + 1, cols=3)
        _style_table(table)
        _set_col_widths(table, [4, 3.5, 9])
        _header_row(table, "Column", "Type", "Notes")
        for i, (col, typ, note) in enumerate(cols):
            _data_row(table, i + 1, col, typ, note, shade=(i % 2 == 1))
            for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
        doc.add_paragraph()


def section_design_decisions(doc):
    add_heading(doc, "10. Key Design Decisions")

    decisions = [
        ("Claude model",           "claude-opus-4-7",
         "Blue chip behavioral context is more nuanced than broad scanning. Higher quality "
         "reasoning on a 10-stock universe justifies the cost."),
        ("3-tier pool",            "Pool 1 → 2 → 3",
         "Behavioral pre-qualification reduces the burden on Claude. Each tier filters for fit, not just technicals."),
        ("Static Pool 2 seed",     "Yes, Phase 1",
         "Need 30+ trading days before Pool Scorer has enough signal to drive dynamic promotion meaningfully."),
        ("No ML scorer",           "Phase 1",
         "Adds complexity without data. Pool behavioral scoring is the equivalent signal."),
        ("Same Supabase",          "b_ prefix",
         "One set of credentials. Combined dashboard. Zero infra overhead."),
        ("Same Alpaca account",    "strategy=b tag",
         "Paper account — no real money distinction needed. Tag enables per-strategy P&L reporting."),
        ("Separate GitHub repo",   "Yes",
         "Zero risk of touching Strategy A code. Independent deployment and failure modes."),
        ("R:R floor = 2.0",        "Lower than A's 3.0",
         "Pool 3 stocks have proven behavioral fit — higher expected hit rate. 2:1 at 75%+ win rate is strongly +EV."),
        ("Intraday target = +1%",  "Capped below premarket",
         "Shorter window means less time for a +2% move to develop. Paired with momentum confirmation, 1.5:1 R:R is still +EV at high win rates."),
        ("INTRADAY_MOMENTUM pool field", "Required in scanner output",
         "Strategy prompt must classify candidates by signal type to apply time-of-day rules. Without pool field, Claude cannot distinguish momentum from regular entries."),
    ]

    table = doc.add_table(rows=len(decisions) + 1, cols=3)
    _style_table(table)
    _set_col_widths(table, [4, 4, 8.5])
    _header_row(table, "Decision", "Choice", "Reason")
    for i, (dec, choice, reason) in enumerate(decisions):
        _data_row(table, i + 1, dec, choice, reason, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY
    doc.add_paragraph()


def section_changelog(doc):
    add_heading(doc, "11. Change Log")

    add_heading(doc, "v1.2 — 2026-05-22", level=2)

    changes = [
        ("Fix",
         "scanner/intraday_momentum.py",
         "Momentum candidate dict now always includes pool=2 alongside signal_type=INTRADAY_MOMENTUM. "
         "Without this field, the strategy prompt could not distinguish intraday momentum candidates "
         "from regular Pool 2 entries and incorrectly applied the 1pm pool restriction."),
        ("Fix",
         "agents/strategy.py (prompt)",
         "Prompt now explicitly exempts signal_type=INTRADAY_MOMENTUM from the 1pm pool restriction. "
         "These are confirmed movers (SPY gate + VWAP confirmation already passed) — they are valid "
         "at any hour, not speculative setups that need time-of-day protection."),
    ]

    table = doc.add_table(rows=len(changes) + 1, cols=3)
    _style_table(table)
    _set_col_widths(table, [2, 5, 9.5])
    _header_row(table, "Type", "File", "Description")
    for i, (typ, file_, desc) in enumerate(changes):
        _data_row(table, i + 1, typ, file_, desc, shade=(i % 2 == 1))
        for run in table.rows[i + 1].cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = ORANGE
    doc.add_paragraph()

    add_heading(doc, "v1.1 — 2026-05-20", level=2)
    for item in [
        "Pool Filter: 5-signal scoring (VWAP, ORB, volume ratio, sector RS, volume acceleration)",
        "Pool Scorer EOD: rolling 7-day scoring, promote/demote logic",
        "Intraday momentum scanner with SPY gate and max-runs guard",
        "Partial profit design (Leg A/Leg B) carried over from Strategy A",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    add_heading(doc, "v1.0 — 2026-05-01", level=2)
    for item in [
        "Initial deployment",
        "Static Pool 2 seed from config",
        "Strategy B premarket pipeline live",
        "b_ prefixed Supabase tables",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    title_block(doc)
    section_overview(doc)
    section_architecture(doc)
    section_pool_system(doc)
    section_agent_inventory(doc)
    section_pipeline(doc)
    section_trading_logic(doc)
    section_risk(doc)
    section_config(doc)
    section_data_model(doc)
    section_design_decisions(doc)
    section_changelog(doc)

    out = "Trading_Agent_B_Design.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
